"""使用BGE-M3语义相似度进行智能分块 - 动态百分位法区分灰色地带，灰色地带交由LLM"""

from langchain_core.documents import Document
from langchain_openai import ChatOpenAI
from FlagEmbedding import BGEM3FlagModel
import numpy as np
import json
import docx
import os
from config import BGE_MODEL_PATH, setup_env
from agent.prompt_loader import load_prompt


setup_env()

_bge_model = None


def _get_bge_model():
    global _bge_model
    if _bge_model is None:
        _bge_model = BGEM3FlagModel(
            model_name_or_path=BGE_MODEL_PATH,
            use_fp16=False,
            device="cpu"
        )
    return _bge_model


def _compute_cosine_similarity(vec1, vec2):
    dot = np.dot(vec1, vec2)
    norm1 = np.linalg.norm(vec1)
    norm2 = np.linalg.norm(vec2)
    if norm1 == 0 or norm2 == 0:
        return 0.0
    return float(dot / (norm1 * norm2))


def _compute_paragraph_similarities(paragraphs):
    if len(paragraphs) <= 1:
        return [], None

    model = _get_bge_model()
    output = model.encode(
        paragraphs,
        return_dense=True,
        return_sparse=False,
        return_colbert_vecs=False
    )
    dense_vecs = np.array(output['dense_vecs'])

    similarities = []
    for i in range(len(dense_vecs) - 1):
        sim = _compute_cosine_similarity(dense_vecs[i], dense_vecs[i + 1])
        similarities.append(sim)

    return similarities, dense_vecs


def _dynamic_percentile_thresholds(similarities, lower_percentile=25, upper_percentile=75):
    if not similarities:
        return 0.0, 0.0

    lower = np.percentile(similarities, lower_percentile)
    upper = np.percentile(similarities, upper_percentile)
    return float(lower), float(upper)


def _classify_boundaries(similarities, lower_threshold, upper_threshold):
    boundaries = []
    for i, sim in enumerate(similarities):
        if sim < lower_threshold:
            boundaries.append(('split', i, sim))
        elif sim > upper_threshold:
            boundaries.append(('merge', i, sim))
        else:
            boundaries.append(('gray', i, sim))
    return boundaries


def _llm_resolve_gray_zones(paragraphs, gray_boundaries, chat_llm):
    prompt = load_prompt("gray_zone_chunker.txt")

    decisions = {}
    for boundary_type, idx, sim in gray_boundaries:
        if boundary_type != 'gray':
            continue
        para_a = paragraphs[idx]
        para_b = paragraphs[idx + 1]

        context = f"段落A:\n{para_a}\n\n段落B:\n{para_b}"

        response = chat_llm.invoke([
            ("system", prompt),
            ("user", context)
        ])

        result = json.loads(response.content)
        decisions[idx] = result.get("should_split", False)

    return decisions


def _build_chunks(paragraphs, boundaries, llm_decisions=None):
    if llm_decisions is None:
        llm_decisions = {}

    split_points = set()
    for boundary_type, idx, _ in boundaries:
        if boundary_type == 'split':
            split_points.add(idx)
        elif boundary_type == 'gray':
            if llm_decisions.get(idx, False):
                split_points.add(idx)

    chunks = []
    start = 0
    for i in range(len(paragraphs)):
        if i in split_points:
            chunk_text = "\n".join(paragraphs[start:i + 1])
            if chunk_text.strip():
                chunks.append(Document(page_content=chunk_text))
            start = i + 1

    if start < len(paragraphs):
        chunk_text = "\n".join(paragraphs[start:])
        if chunk_text.strip():
            chunks.append(Document(page_content=chunk_text))

    return chunks


def load_word_document_file(file_path, chat_llm=None, lower_percentile=25, upper_percentile=75):
    abs_path = os.path.abspath(file_path)
    doc = docx.Document(abs_path)
    texts = []
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text)

    print(f"成功加载Word文档，共解析出 {len(texts)} 个段落")

    if len(texts) <= 1:
        return [Document(page_content=texts[0])] if texts else []

    similarities, _ = _compute_paragraph_similarities(texts)

    lower, upper = _dynamic_percentile_thresholds(similarities, lower_percentile, upper_percentile)
    print(f"动态阈值: 下界(P{lower_percentile})={lower:.3f}, 上界(P{upper_percentile})={upper:.3f}")

    boundaries = _classify_boundaries(similarities, lower, upper)
    gray_count = sum(1 for t, _, _ in boundaries if t == 'gray')
    split_count = sum(1 for t, _, _ in boundaries if t == 'split')
    merge_count = sum(1 for t, _, _ in boundaries if t == 'merge')
    print(f"边界分类: 明确分割={split_count}, 明确合并={merge_count}, 灰色地带={gray_count}")

    llm_decisions = {}
    if gray_count > 0 and chat_llm is not None:
        print(f"使用LLM处理 {gray_count} 个灰色地带...")
        gray_boundaries = [(t, i, s) for t, i, s in boundaries if t == 'gray']
        llm_decisions = _llm_resolve_gray_zones(texts, gray_boundaries, chat_llm)
        llm_split = sum(1 for v in llm_decisions.values() if v)
        llm_merge = sum(1 for v in llm_decisions.values() if not v)
        print(f"LLM判断: {llm_split} 个分割, {llm_merge} 个合并")
    elif gray_count > 0:
        print(f"有 {gray_count} 个灰色地带但未提供LLM，默认合并处理")

    chunks = _build_chunks(texts, boundaries, llm_decisions)
    print(f"分块后共 {len(chunks)} 个chunk")
    return chunks
